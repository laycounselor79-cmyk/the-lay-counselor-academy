#!/usr/bin/env python3
"""Build styled .docx files from SITE_COPY_FOR_REVIEW.md.

Produces:
  - SITE_COPY_FOR_REVIEW.docx                (single combined doc)
  - review-tabs/00-read-me-first.docx        (intro / how-to-edit)
  - review-tabs/01-home.docx ... 11-404.docx (one per page, for Google Docs tabs)
"""
import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "SITE_COPY_FOR_REVIEW.md"
COMBINED_OUT = ROOT / "SITE_COPY_FOR_REVIEW.docx"
TABS_DIR = ROOT / "review-tabs"

NAVY = RGBColor(0x14, 0x2C, 0x53)
CORAL = RGBColor(0xC8, 0x4A, 0x3D)
INK = RGBColor(0x1C, 0x22, 0x2E)
INK_SOFT = RGBColor(0x4A, 0x52, 0x60)
INK_QUIET = RGBColor(0x6E, 0x76, 0x83)

SERIF = "Georgia"
SANS = "Helvetica Neue"


def set_run(run, *, font=None, size=None, bold=None, italic=None, color=None):
    if font:
        run.font.name = font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_bottom_border(paragraph, color="D9D3C8", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_inline_runs(paragraph, text, *, base_font=SANS, base_size=11, base_color=INK):
    token_re = re.compile(
        r"(\*\*[^*]+\*\*)"
        r"|(\*[^*]+\*)"
        r"|(\[[^\]]+\]\([^)]+\))"
        r"|(\[[A-Z0-9][^\]]*\])"
    )
    pos = 0
    for m in token_re.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_run(r, font=base_font, size=base_size, color=base_color)
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            set_run(r, font=base_font, size=base_size, color=base_color, bold=True)
        elif token.startswith("*"):
            r = paragraph.add_run(token[1:-1])
            set_run(r, font=base_font, size=base_size, color=base_color, italic=True)
        elif token.startswith("[") and "](" in token:
            close = token.index("]")
            label = token[1:close]
            url = token[close + 2:-1]
            r = paragraph.add_run(label)
            set_run(r, font=base_font, size=base_size, color=RGBColor(0x1F, 0x49, 0x8E))
            r.font.underline = True
            r2 = paragraph.add_run(f" ({url})")
            set_run(r2, font=base_font, size=base_size - 1, color=INK_QUIET)
        elif token.startswith("["):
            r = paragraph.add_run(token)
            set_run(r, font=SANS, size=base_size - 1, color=CORAL, bold=True)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run(r, font=base_font, size=base_size, color=base_color)


def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = SANS
    style.font.size = Pt(11)
    style.font.color.rgb = INK
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
    return doc


def render(doc, lines):
    i = 0
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        rows = [row for row in table_rows if not all(set(c.strip()) <= set("-:") for c in row)]
        if not rows:
            table_rows = []
            return
        cols = len(rows[0])
        tbl = doc.add_table(rows=len(rows), cols=cols)
        tbl.style = "Light Grid Accent 1"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                cell = tbl.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                add_inline_runs(
                    p, cell_text.strip(),
                    base_font=SANS, base_size=10,
                    base_color=NAVY if r_idx == 0 else INK,
                )
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        elif in_table:
            flush_table()
            in_table = False

        if not stripped:
            i += 1
            continue

        if stripped.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            add_bottom_border(p, color="D9D3C8", size=8)
            i += 1
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(28)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.keep_with_next = True
            add_bottom_border(p, color="142C53", size=12)
            r = p.add_run(text)
            set_run(r, font=SERIF, size=24, bold=False, color=NAVY)
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(text)
            set_run(r, font=SERIF, size=16, bold=False, color=NAVY)
            i += 1
            continue

        if stripped.startswith("### "):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(text)
            set_run(r, font=SERIF, size=13, bold=True, color=NAVY)
            i += 1
            continue

        if re.match(r"^\s*[-*]\s+", stripped):
            text = re.sub(r"^\s*[-*]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_inline_runs(p, text, base_font=SANS, base_size=11, base_color=INK_SOFT)
            i += 1
            continue

        if stripped.startswith("*(") and stripped.endswith(")*"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(stripped[1:-1])
            set_run(r, font=SERIF, size=10, italic=True, color=INK_QUIET)
            i += 1
            continue

        if stripped == "*End of document.*":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            r = p.add_run("End of document.")
            set_run(r, font=SERIF, size=10, italic=True, color=INK_QUIET)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.35
        add_inline_runs(p, stripped, base_font=SANS, base_size=11, base_color=INK_SOFT)
        i += 1

    if in_table:
        flush_table()


def split_into_sections(md_text):
    """Split markdown into intro + per-page sections.

    A page boundary is a top-level `# ` heading whose text starts with a digit
    (e.g. `# 1. HOME PAGE`). Other `# ` headings (the doc title) stay in intro.
    """
    lines = md_text.split("\n")
    intro = []
    sections = []
    current_heading = None
    current_body = []
    for line in lines:
        is_page_heading = (
            line.startswith("# ")
            and not line.startswith("## ")
            and re.match(r"^# \d", line)
        )
        if is_page_heading:
            if current_heading is not None:
                sections.append((current_heading, current_body))
            else:
                intro = current_body
            current_heading = line[2:].strip()
            current_body = []
        else:
            current_body.append(line)
    if current_heading is not None:
        sections.append((current_heading, current_body))
    else:
        intro = current_body
    return intro, sections


def slugify(heading):
    """Turn '1. HOME PAGE  (`/`)' into 'home', '7. ELIZABETH MORRISON BIO ...' into 'elizabeth-morrison'."""
    # Strip leading number + dot
    text = re.sub(r"^\d+\.\s*", "", heading)
    # Drop trailing parenthetical path
    text = re.sub(r"\s*\(.*$", "", text).strip().lower()
    # Drop the word "page"
    text = re.sub(r"\bpage\b", "", text).strip()
    # Drop "bio"
    text = re.sub(r"\bbio\b", "", text).strip()
    # Collapse to slug
    text = re.sub(r"[^a-z0-9]+", "-", text).strip("-")
    return text or "section"


def build_combined(md_text):
    doc = new_doc()
    render(doc, md_text.split("\n"))
    doc.save(COMBINED_OUT)
    print(f"Wrote {COMBINED_OUT.relative_to(ROOT)} ({COMBINED_OUT.stat().st_size:,} bytes)")


def build_tabs(md_text):
    if TABS_DIR.exists():
        shutil.rmtree(TABS_DIR)
    TABS_DIR.mkdir()

    intro_lines, sections = split_into_sections(md_text)

    # Trim trailing whitespace/end-of-doc marker from intro
    while intro_lines and not intro_lines[-1].strip():
        intro_lines.pop()

    # 00 — read-me-first tab (the intro / how-to-edit guidance)
    doc = new_doc()
    render(doc, ["# Read me first"] + [""] + intro_lines)
    path = TABS_DIR / "00-read-me-first.docx"
    doc.save(path)
    print(f"Wrote {path.relative_to(ROOT)}")

    # One file per page
    for idx, (heading, body) in enumerate(sections, start=1):
        slug = slugify(heading)
        doc = new_doc()
        # Re-emit the H1 heading at the top so the tab content is self-contained
        render(doc, [f"# {heading}"] + body)
        path = TABS_DIR / f"{idx:02d}-{slug}.docx"
        doc.save(path)
        print(f"Wrote {path.relative_to(ROOT)}  ({heading})")


def main():
    md = SRC.read_text()
    build_combined(md)
    build_tabs(md)


if __name__ == "__main__":
    main()
